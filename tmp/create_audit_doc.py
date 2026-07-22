from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path('output/documents/cinema-booking-requirements-audit.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = '0B214A'
BLUE = '2E74B5'
LIGHT_BLUE = 'E8EEF5'
PALE_GREEN = 'EAF4ED'
PALE_AMBER = 'FFF4DF'
PALE_RED = 'FCEAEA'
TEXT = '1F2937'
MUTED = '5B6776'
WHITE = 'FFFFFF'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')

def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in('w:tblLayout')
    if tbl_layout is None:
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn('w:type'), 'fixed')
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)

def set_font(run, size=None, color=TEXT, bold=None, italic=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    if size:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def para(doc, text='', style=None, after=6, before=0, color=TEXT, bold=False, size=None, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=11)
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    r = p.add_run(text)
    set_font(r, size={1:16, 2:13, 3:12}[level], color=BLUE if level < 3 else NAVY, bold=True)
    return p

def add_table(doc, headers, rows, widths, status_column=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_geometry(table, widths)
    header = table.rows[0]
    repeat_header(header)
    for i, label in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(label)
        set_font(r, size=9.5, color=WHITE, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell = cells[i]
            fill = None
            if status_column is not None and i == status_column:
                lowered = str(value).lower()
                fill = PALE_GREEN if 'done' in lowered or 'implemented' in lowered else PALE_AMBER if 'partial' in lowered else PALE_RED
            if fill:
                set_cell_shading(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(str(value))
            set_font(r, size=9.2, color=TEXT, bold=(status_column == i))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, size=10.5, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(text)
    set_font(r2, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for level, size, before, after in [(1,16,18,10),(2,13,14,7),(3,12,10,5)]:
    style = styles[f'Heading {level}']
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else NAVY)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = header.add_run('BEATRIX MOVIE | REQUIREMENTS AUDIT')
set_font(run, size=8.5, color=MUTED, bold=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Current synchronized codebase | July 2026 | Page ')
set_font(run, size=8, color=MUTED)
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'PAGE')
footer._p.append(fld)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(42)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('TECHNICAL AUDIT')
set_font(r, size=11, color=BLUE, bold=True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run('Cinema Booking System\nRequirement Coverage & API Reference')
set_font(r, size=28, color=NAVY, bold=True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(24)
r = p.add_run('Assessment of the current synchronized repository against the supplied Full-Stack Development Challenge requirements.')
set_font(r, size=13, color=MUTED)

meta = doc.add_table(rows=3, cols=2)
set_table_geometry(meta, [2700, 6660])
for i, (label, value) in enumerate([
    ('Audit basis', 'Challenge requirements PDF plus static inspection of frontend and backend code'),
    ('Scope', 'Implemented work, remaining requirements, actual API routes, and repository documentation'),
    ('Assessment type', 'Static code review; MongoDB/runtime behavior was not executed as part of this report')
]):
    set_cell_shading(meta.cell(i,0), LIGHT_BLUE)
    rp = meta.cell(i,0).paragraphs[0].add_run(label)
    set_font(rp, size=10, color=NAVY, bold=True)
    rv = meta.cell(i,1).paragraphs[0].add_run(value)
    set_font(rv, size=10)

doc.add_paragraph()
callout(doc, 'Assessment conclusion', 'The project has a strong UI and module foundation, but it is not yet compliant with the mandatory secure booking flow. The highest-risk gaps are unprotected admin and write routes, missing booking ownership/showtime references, and incomplete seat reservation/cancellation behavior.', PALE_AMBER)

heading(doc, '1. Executive compliance matrix', 1)
para(doc, 'Status definitions: Implemented = present and substantially wired; Partial = present but not connected or not requirement-safe; Not implemented = required flow is absent.', after=8, color=MUTED, size=10)
add_table(doc,
    ['Requirement area', 'Status', 'Evidence / assessment'],
    [
        ('React conversion and component structure', 'Implemented', 'React/Vite pages, shared layout, feature folders, and stateful booking UI are present.'),
        ('Movie catalogue, search, filters, pagination', 'Implemented', 'GET /api/movies supports search/q, genre, status, page, limit, and sort.'),
        ('Showtime list and seat map', 'Partial', 'API-backed UI exists, but local fallback data remains and booked seats are not updated from checkout.'),
        ('JWT registration/login/current user/logout', 'Partial', 'Controller and middleware exist, but backend dependency manifest omits bcryptjs, jsonwebtoken, and cookie-parser.'),
        ('Frontend login/register and route guards', 'Not implemented', 'Pages and guard components exist but are not mounted by App.jsx hash routing.'),
        ('Admin frontend and protected admin access', 'Partial', 'Admin UI exists, but #admin is unguarded and booking/stat data is mocked.'),
        ('Admin API enforcement', 'Not implemented', 'Admin routes are not mounted; movie/showtime write endpoints are public.'),
        ('Booking ownership and My Bookings', 'Not implemented', 'Booking schema has no user reference and no /api/bookings/me route.'),
        ('Seat conflict prevention per showtime', 'Partial', 'A basic booking conflict query exists, but it is not tied to Showtime.bookedSeats or showtimeId.'),
        ('Cancellation and seat release', 'Not implemented', 'No owner/admin cancellation route or release logic.'),
        ('Seed data and Postman collection', 'Not implemented', 'Movies/showtimes auto-seed only; users/admin, seed command, and exported collection are absent.'),
        ('Documentation', 'Partial', 'Useful docs exist, but API contract/readmes describe planned or placeholder routes and an older folder tree.')
    ], [3000, 1500, 4860], status_column=1)

heading(doc, '2. Implemented work', 1)
heading(doc, '2.1 Frontend and user experience', 2)
for item in [
    'Home, movie-detail, showtime selection, seat map, payment, ticket, and admin screens are implemented as React components.',
    'Movie browsing retrieves backend data when available and supports search, genre/status filters, sorting, and pagination.',
    'Seat UI differentiates available, selected, and booked states; showtime seat availability is exposed through a dedicated endpoint.',
    'Mock QRIS checkout, payment-success simulation, ticket QR code, ticket PDF generation, and mock email preview are implemented as bonus features.'
]: bullet(doc, item)
heading(doc, '2.2 Backend foundation', 2)
for item in [
    'Express server with CORS, JSON parsing, cookies, MongoDB connection, modular route/controller/model organization, and central error middleware.',
    'Mongoose models for users, movies, showtimes, and bookings; Showtime has a movie ObjectId relationship and validates seat IDs.',
    'Authentication controller contains registration, login, HTTP-only JWT cookie, current-user, and logout logic.',
    'Movie and showtime controllers provide read operations, validation, CRUD methods, and movie-list query parameters.'
]: bullet(doc, item)

heading(doc, '3. Required work not yet complete', 1)
callout(doc, 'Priority 1 - security and integrity', 'Do not treat the current QRIS/ticket flow as challenge-complete booking logic. It remains public and does not bind a booking to the authenticated user or selected showtime.', PALE_RED)
for item in [
    'Add backend dependencies: bcryptjs, jsonwebtoken, and cookie-parser. A clean npm install currently cannot start the imported authentication code.',
    'Mount /api/admin routes in server.js. Protect admin stats/bookings and movie/showtime writes with authenticate + requireAdmin.',
    'Make the frontend use one route system. Mount LoginPage, RegisterPage, ProtectedRoute, and AdminRoute; protect seat selection, My Bookings, and admin pages.',
    'Redesign Booking around user ObjectId and showtime ObjectId. Derive price from Showtime.price, not Movie.price or client input.',
    'Atomically re-check and reserve seats in Showtime.bookedSeats. On conflict, return HTTP 409 with unavailableSeats and refresh the map.',
    'Add owner/admin-safe booking detail, My Bookings, cancellation, and seat-release flows.',
    'Replace admin sample booking/stat cards with protected API responses. Implement stats and all-bookings controllers instead of 501 placeholders.',
    'Add a documented seed command that creates one admin, two users, movies, and multiple showtimes; add an exported Postman collection.',
    'Return safe error messages without development stack traces and update all documentation to describe the active system.'
]: bullet(doc, item)

heading(doc, '4. Actual API reference', 1)
para(doc, 'The following is the effective current route surface. “Public” means no authentication middleware is attached in the current router.', after=8, color=MUTED, size=10)
heading(doc, '4.1 Authentication', 2)
add_table(doc, ['Method', 'Route', 'Effective access', 'Purpose'], [
    ('POST', '/api/auth/register', 'Public', 'Register normal user; role is forced to user.'),
    ('POST', '/api/auth/login', 'Public', 'Verify credentials and set HTTP-only JWT cookie.'),
    ('POST', '/api/auth/logout', 'Authenticated', 'Clear session cookie.'),
    ('GET', '/api/auth/me', 'Authenticated', 'Return authenticated user profile and role.')
], [1100, 2900, 1800, 3560])
heading(doc, '4.2 Movies and showtimes', 2)
add_table(doc, ['Method', 'Route', 'Effective access', 'Purpose'], [
    ('GET', '/api/movies', 'Public', 'List movies; supports search/q, genre, status, page, limit, sort.'),
    ('GET', '/api/movies/genres', 'Public', 'List available genres.'),
    ('GET', '/api/movies/:idOrSlug', 'Public', 'Get movie detail.'),
    ('GET', '/api/movies/:movieId/showtimes', 'Public', 'List showtimes for one movie.'),
    ('POST/PUT/DELETE', '/api/movies...', 'Public - should be admin', 'Create, update, delete movie.'),
    ('GET', '/api/showtimes', 'Public', 'List showtimes; supports movieId/date/from/studio/page/limit.'),
    ('GET', '/api/showtimes/:id', 'Public', 'Get showtime detail.'),
    ('GET', '/api/showtimes/:id/seats', 'Public', 'Return layout and latest showtime booked seats.'),
    ('POST/PUT/DELETE', '/api/showtimes...', 'Public - should be admin', 'Create, update, delete showtime.')
], [1200, 2800, 2100, 3260])
heading(doc, '4.3 Bookings, payment, and tickets', 2)
add_table(doc, ['Method', 'Route', 'Effective access', 'Purpose'], [
    ('GET', '/api/bookings', 'Public', 'Lists all bookings; does not enforce ownership.'),
    ('POST', '/api/bookings', 'Public', 'Creates raw booking document; not challenge-safe.'),
    ('GET', '/api/bookings/occupied', 'Public', 'Basic occupied-seat lookup using movie/cinema/showtime query.'),
    ('POST', '/api/bookings/checkout', 'Public', 'Create mock QRIS checkout and pending booking.'),
    ('PATCH', '/api/bookings/:id/mock-paid', 'Public', 'Simulate payment success and issue ticket.'),
    ('POST', '/api/bookings/:id/email', 'Public', 'Send mock ticket email.'),
    ('GET', '/api/bookings/:id/ticket.pdf', 'Public', 'Download ticket PDF after payment.')
], [1200, 3000, 1900, 3260])
heading(doc, '4.4 Admin routes', 2)
add_table(doc, ['Method', 'Route', 'Runtime status', 'Notes'], [
    ('GET', '/api/admin/stats', 'Not mounted', 'Route file exists; controller currently returns 501 if mounted.'),
    ('GET', '/api/admin/bookings', 'Not mounted', 'Route file exists; controller currently returns 501 if mounted.')
], [1200, 3000, 1900, 3260])

heading(doc, '5. Required API additions and access rules', 1)
add_table(doc, ['Method', 'Required route', 'Required access', 'Needed behavior'], [
    ('GET', '/api/bookings/me', 'Authenticated', 'Return only bookings owned by req.user.'),
    ('GET', '/api/bookings/:id', 'Owner or admin', 'Return booking only after ownership/role check.'),
    ('DELETE', '/api/bookings/:id', 'Owner or admin', 'Cancel booking and release seats on the related showtime.'),
    ('POST', '/api/bookings', 'Authenticated', 'Validate showtime and seats; reserve current seats atomically.'),
    ('GET', '/api/admin/stats', 'Admin', 'Return counts for users, movies, showtimes, and bookings.'),
    ('GET', '/api/admin/bookings', 'Admin', 'Return all bookings with populated user/movie/showtime data.')
], [1100, 3000, 1900, 3360])

heading(doc, '6. Current project structure', 1)
para(doc, 'The top-level split is appropriate: frontend/ contains React/Vite code and backend/ contains Express/Mongoose code.', after=6)
para(doc, 'Current high-level structure:', after=4, color=MUTED, size=10)
tree = '''frontend/
  src/
    app/                 App shell and hash-route metadata
    pages/               Home, movie detail, seat selection, payment, admin, auth pages
    features/            movies, showtimes, bookings, auth, admin
    components/layout/   Header and footer
    services/            Generic API client
    hooks/ and utils/    Shared client helpers
backend/
  src/
    config/              MongoDB connection
    modules/             auth, movies, showtimes, bookings, admin
    shared/              services, utilities, middleware placeholders
    data/                sample movies and showtimes
    server.js            Express composition and route mounting'''
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
for line in tree.splitlines():
    r = p.add_run(line + '\n')
    set_font(r, size=9.4, color=TEXT)
    r.font.name = 'Consolas'
    r._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Consolas')
heading(doc, '6.1 Existing documentation', 2)
add_table(doc, ['Document', 'Current value', 'Required update'], [
    ('README.md', 'Explains two-folder project and local run commands.', 'Replace the old shared/services folder tree with the current route-first tree.'),
    ('API_CONTRACT.md', 'Best endpoint reference; distinguishes active and planned work.', 'Document the QRIS/ticket endpoints and correct effective access status.'),
    ('backend/README.md', 'Lists backend modules and route inventory.', 'Remove “placeholder” wording once implemented; do not list unmounted admin routes as available.'),
    ('RUNNING_GUIDE.md', 'Explains install, env, and basic demo flow.', 'Add authentication, seed credentials, payment/ticket flow, and required test scenarios.'),
    ('PROJECT_STRUCTURE.md / module READMEs', 'Describe organization and ownership boundaries.', 'Keep aligned with current imports and completed feature status.')
], [1900, 3700, 3760])

heading(doc, '7. Recommended remediation sequence', 1)
for number, text in enumerate([
    'Repair backend dependency manifest and verify authentication startup from a clean install.',
    'Mount the admin router and add authentication/administrator middleware to all protected operations.',
    'Wire login/register and protected route behavior into the live frontend route system.',
    'Implement authoritative booking ownership and atomic showtime seat reservation.',
    'Implement My Bookings, cancellation, and seat release; then replace admin mock data with real APIs.',
    'Add seed users/admin, Postman test collection, and documentation updates; demonstrate the required scenarios.'
], 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=11)

callout(doc, 'Definition of done', 'The requirements are satisfied when a registered user can authenticate, select a showtime, reserve seats through protected backend validation, view/cancel only their bookings, and when an administrator can access protected CRUD and monitoring APIs. The current mock QRIS/ticket flow can remain as a bonus once that core path is secured.', PALE_GREEN)

doc.core_properties.title = 'Cinema Booking System Requirements Audit'
doc.core_properties.subject = 'Requirement coverage, API reference, and documentation review'
doc.core_properties.author = 'Beatrix Movie Team'
doc.save(OUT)
print(OUT)
